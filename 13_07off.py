import streamlit as st
import os
from PyPDF2 import PdfReader
from langchain.chains import RetrievalQA
from langchain.chat_models import ChatOpenAI
from langchain_google_genai import ChatGoogleGenerativeAI
import string 
import docx
from langchain.embeddings import HuggingFaceBgeEmbeddings
from langchain.vectorstores import Qdrant
import random
from datetime import datetime
from langchain.text_splitter import CharacterTextSplitter
from qdrant_client import QdrantClient
from streamlit_chat import message
from langchain.docstore.document import Document
from dotenv import load_dotenv

openai_api = st.secrets["OPENAI_API_KEY"]
qdrant_api = st.secrets["QDRANT_API_KEY"]
qdrant_url = st.secrets["QDRANT_URL"]
google_api_key = st.secrets["GOOGLE_API_KEY"]
embaddings = HuggingFaceBgeEmbeddings(model_name = "sentence-transformers/all-mpnet-base-v2")

def main():
    load_dotenv()
    st.set_page_config('Document GPT')
    st.title("Document GPT with Reference")

    if 'conversation' not in st.session_state:
        st.session_state.conversation = []
    if 'chat_history' not in st.session_state:
        st.session_state.chat_history = []
    if 'processComplete' not in st.session_state:
        st.session_state.processComplete = None

    with st.sidebar:
        uploaded_files = st.file_uploader("Please upload your file!", type=['pdf'], accept_multiple_files=True)
        llm_model = st.selectbox("Please Select your desired model! ", options=['Gemini', "ChatGPT"])
        if llm_model == 'Gemini':
            llm_model= ChatGoogleGenerativeAI(model= "gemini-1.5-flash")
        elif llm_model == 'ChatGPT':
            llm_model = ChatOpenAI(model = "gpt-3.5-turbo")
        process = st.button('Process')
    if process:
        if not openai_api or not google_api_key:
            st.error("Please set OPENAI_API_KEY in .env file")
            st.stop()
        text_chunks_list = []
        for uploaded_file in uploaded_files:
            file_name = uploaded_file.name
            file_text = get_files_text(uploaded_file)
            text_chunks = get_text_chunks(file_text, file_name)
            text_chunks_list.extend(text_chunks)
            
        curr_date = str(datetime.now())
        collection_name = "".join(random.choices(string.ascii_letters, k = 4)) + curr_date.split('.')[0].replace(':', '-').replace(" ", 'T')
        vectorestore = get_vectorstore(text_chunks_list, collection_name)
        st.write('Vector Store Created...')
        num_chunks = 4
        st.session_state.conversation = get_qa_chain(vectorestore, num_chunks, llm_model)
        st.session_state.processComplete = True
    if st.session_state.processComplete == True:
        user_question = st.chat_input("Please add your question ")
        if user_question:
            handel_userinput(user_question)


def get_files_text(uploaded_file):
    text = ""
    split_tup = os.path.splitext(uploaded_file.name)
    file_extension = split_tup[1]
    if file_extension == ".pdf":
        text += get_pdf_text(uploaded_file)
    elif file_extension == ".docx":
        text += get_docx_text(uploaded_file)
    else:
        pass
    return text

# Function to read PDF Files
def get_pdf_text(pdf):
    pdf_reader = PdfReader(pdf)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

def get_docx_text(file):
    doc = docx.Document(file)
    allText = []
    for docpara in doc.paragraphs:
        allText.append(docpara.text)
    text = ' '.join(allText)
    return text



def get_text_chunks(text, filename):
    # spilit ito chuncks
    text_splitter = CharacterTextSplitter(
        separator="\n",
        chunk_size=80,
        chunk_overlap=20,
        length_function=len
    )
    chunks = text_splitter.split_text(text)
    doc_list = []
    for chunk in chunks:
        metadata = {"source": filename}
        doc_string = Document(page_content=chunk, metadata=metadata)
        doc_list.append(doc_string)
    return doc_list


def get_vectorstore(text_chunks, COLLECTION_NAME):
    # Using the hugging face embedding models
    try:
        # creating the Vectore Store using Facebook AI Semantic search
        knowledge_base = Qdrant.from_documents(
            documents = text_chunks,
            embedding = embaddings,
            url=qdrant_url,
            prefer_grpc=True,
            api_key=qdrant_api,
            collection_name=COLLECTION_NAME,
        )
    except Exception as e:
        st.write(f"Error: {e}")
    return knowledge_base

def get_qa_chain(vectorstore,num_chunks, llm_model):
    # prompt_template = """
    # You are trained to extract Answer from the given Context and Question. Then, precise the Answer in less than 20 words. If the Answer is not found in the Context, then return "N/A", otherwise return the precise Answer.
    # Context: {context}
    # Question: {question}"""
    # mprompt_url = PromptTemplate(
    #     template=prompt_template, input_variables=["context", "question"], validate_template=False)
    # chain_type_kwargs = {"prompt": mprompt_url}


    # qa = RetrievalQA.from_chain_type(llm=ChatOpenAI(model = "gpt-3.5-turbo"), chain_type="stuff",
    #                             retriever=vectorstore.as_retriever(search_type="similarity",
    #                                                         search_kwargs={"k": num_chunks}), chain_type_kwargs=chain_type_kwargs, return_source_documents=True)
    qa = RetrievalQA.from_chain_type(llm=llm_model, chain_type="stuff",
                                retriever=vectorstore.as_retriever(search_type="similarity",
                                                            search_kwargs={"k": num_chunks}),  return_source_documents=True)
    return qa


def handel_userinput(user_question):
    with st.spinner('Generating response...'):
        result = st.session_state.conversation({"query": user_question})
        response = result['result']
        source = result['source_documents'][0].metadata['source']
    st.session_state.chat_history.append(user_question)
    st.session_state.chat_history.append(f"{response} \n Source Document: {source}")


    # Layout of input/response containers
    response_container = st.container()

    with response_container:
        for i, messages in enumerate(st.session_state.chat_history):
            if i % 2 == 0:
                message(messages, is_user=True, key=str(i))
            else:
                message(messages, key=str(i))


if __name__ == '__main__':
    main()





